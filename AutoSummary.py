# AutoSummary.py
import sys
import os
import json
import subprocess
from pathlib import Path
from datetime import datetime

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout,
    QHBoxLayout, QGridLayout, QPushButton, QLabel, QMessageBox
)
from PyQt5.QtGui import QFont
from PyQt5.QtCore import Qt

from docx import Document
try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except Exception:
    DOCX2PDF_AVAILABLE = False


BASE_DIR = Path(__file__).parent.resolve()
PRODUCT_DIR = BASE_DIR / "Product"
PRODUCT_DIR.mkdir(parents=True, exist_ok=True)
SITES_JSON = BASE_DIR / "sites.json"

MONTHS_FULL = [
    "January","February","March","April","May","June",
    "July","August","September","October","November","December"
]
MONTH_TO_NUM = {m: i+1 for i, m in enumerate(MONTHS_FULL)}
NUM_TO_MONTH = {i+1: m for i, m in enumerate(MONTHS_FULL)}

def load_sites():
    if not SITES_JSON.exists():
        raise FileNotFoundError(f"Missing sites.json at {SITES_JSON}")
    with open(SITES_JSON, "r", encoding="utf-8") as f:
        data = json.load(f)
    # Add excel filename guess if not provided
    for s in data:
        if "excel" not in s or not s["excel"]:
            s["excel"] = f"{s['site']}.xlsx"
    return data

def find_excel_for(site_entry: dict) -> Path | None:
    """Prefer explicit filename; otherwise try to find by stem or by site name."""
    explicit = BASE_DIR / site_entry["excel"]
    if explicit.exists():
        return explicit
    stem = Path(site_entry["excel"]).stem.lower()
    for p in BASE_DIR.glob("*.xls*"):
        if p.stem.lower() == stem or p.stem.lower() == site_entry["site"].lower():
            return p
    return None

def months_between_inclusive(prev_month: str, curr_month: str) -> list[str]:
    """Calendar walk from prev→curr inclusive, wrapping year if needed."""
    start = MONTH_TO_NUM[prev_month]
    end = MONTH_TO_NUM[curr_month]
    out = []
    i = start
    while True:
        out.append(NUM_TO_MONTH[i])
        if i == end:
            break
        i = 1 if i == 12 else i + 1
    return out

def previous_visit_month(visit_list: list[str], current_month: str) -> str | None:
    """Return the previous entry in the ordered visit_list relative to current_month."""
    try:
        idx = visit_list.index(current_month)
        return visit_list[idx-1] if idx > 0 else visit_list[-1]
    except ValueError:
        return None

class SummaryGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AutoSummary")
        self.resize(600, 350)
        self.selected_year = None
        self.selected_month = None

        root = QWidget()
        self.setCentralWidget(root)
        v = QVBoxLayout(root)

        title = QLabel("AutoSummary")
        title.setAlignment(Qt.AlignCenter)
        title.setFont(QFont("Courier", 16, QFont.Bold))
        v.addWidget(title)

        # Month buttons
        grid = QGridLayout()
        self.month_btns = []
        short_months = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
        for i, m in enumerate(short_months):
            btn = QPushButton(m)
            btn.setFixedSize(90, 40)
            btn.clicked.connect(lambda _, idx=i+1, b=btn: self.pick_month(idx, b))
            self.month_btns.append(btn)
            r, c = divmod(i, 6)
            grid.addWidget(btn, r, c)
        v.addLayout(grid)

        # Year buttons (tweak as you like)
        yh = QHBoxLayout()
        self.year_btns = []
        for y in [2023, 2024, 2025]:
            b = QPushButton(str(y))
            b.setFixedSize(100, 36)
            b.clicked.connect(lambda _, year=y, btn=b: self.pick_year(year, btn))
            self.year_btns.append(b)
            yh.addWidget(b)
        v.addLayout(yh)

        # Go button
        go = QPushButton("Create Summaries")
        go.setFixedHeight(48)
        go.clicked.connect(self.run_summary)
        v.addWidget(go)

    def pick_month(self, mnum: int, btn: QPushButton):
        self.selected_month = mnum
        for b in self.month_btns:
            b.setStyleSheet("")
        btn.setStyleSheet("background:#7FDBFF;")

    def pick_year(self, year: int, btn: QPushButton):
        self.selected_year = year
        for b in self.year_btns:
            b.setStyleSheet("")
        btn.setStyleSheet("background:#7FDBFF;")

    def run_summary(self):
        if not self.selected_month or not self.selected_year:
            QMessageBox.warning(self, "Missing selection", "Pick a month and a year.")
            return

        try:
            sites = load_sites()
        except Exception as e:
            QMessageBox.critical(self, "sites.json error", str(e))
            return

        picked_month_name = NUM_TO_MONTH[self.selected_month]
        year = self.selected_year

        processed_any = False

        for site_entry in sites:
            visits = site_entry.get("months", [])
            if picked_month_name not in visits:
                continue

            excel_path = find_excel_for(site_entry)
            if not excel_path:
                print(f"[WARN] Excel not found for site '{site_entry['site']}'")
                continue

            processed_any = True
            person = (site_entry.get("person") or "Unassigned").strip()

            # Technician + month subfolder: Product/<Tech> - <Month> <Year>/
            tech_folder_name = f"{person} - {picked_month_name} {year}"
            TECH_DIR = PRODUCT_DIR / tech_folder_name
            TECH_DIR.mkdir(parents=True, exist_ok=True)

            # Doc path per site inside that folder
            out_name = f"Summary - {site_entry['site']} - {picked_month_name} {year}.docx"
            out_path = TECH_DIR / out_name

            # Seed document
            doc = Document()
            doc.add_heading(site_entry["site"], level=0)
            doc.add_paragraph(f"Assigned to: {person}")
            doc.add_paragraph(f"Reporting Month: {picked_month_name} {year}")
            doc.add_paragraph("")
            doc.add_heading("Results and Discussion", level=1)
            doc.save(out_path)

            # Determine the month span (previous visit -> selected), inclusive
            prev_visit = previous_visit_month(visits, picked_month_name)
            if prev_visit:
                span_months = months_between_inclusive(prev_visit, picked_month_name)
            else:
                span_months = [picked_month_name]

            month_numbers = [str(MONTH_TO_NUM[m]) for m in span_months]
            months_csv = ",".join(month_numbers)

            # --- SGSData: now takes the SAME months_csv window as FlowData ---
            subprocess.run(
                [sys.executable, str(BASE_DIR / "SGSData.py"),
                 str(excel_path), months_csv, str(year), str(out_path)],
                check=False
            )

            # --- FlowData: unchanged; uses the same months_csv window ---
            subprocess.run(
                [sys.executable, str(BASE_DIR / "FlowData.py"),
                 str(excel_path), months_csv, str(year), str(out_path)],
                check=False
            )

            print(f"✔ Wrote: {out_path}")
            export_docx_to_pdf(out_path)


        if not processed_any:
            QMessageBox.information(self, "Nothing to do",
                                    f"No sites scheduled for {picked_month_name}.")
        else:
            QMessageBox.information(
                self, "Done",
                "Summaries created in Product/<Technician - Month Year>/"
            )
def export_docx_to_pdf(docx_path: Path):
    """
    Export a DOCX to a PDF in the same folder. Prefers docx2pdf (Word-based).
    Falls back to Word COM automation if docx2pdf is unavailable.
    """
    pdf_path = docx_path.with_suffix(".pdf")

    # Preferred: docx2pdf (fast + robust)
    if DOCX2PDF_AVAILABLE:
        try:
            docx2pdf_convert(str(docx_path), str(pdf_path))
            print(f"  → PDF saved: {pdf_path}")
            return
        except Exception as e:
            print(f"[WARN] docx2pdf failed: {e}. Trying COM automation…")

    # Fallback: COM automation (Windows + Word required)
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path))
        # 17 = wdFormatPDF
        doc.SaveAs(str(pdf_path), FileFormat=17)
        doc.Close(False)
        word.Quit()
        print(f"  → PDF saved: {pdf_path}")
    except Exception as e:
        print(f"[WARN] PDF export failed: {e}. You can still open the DOCX.")

def main():
    app = QApplication(sys.argv)
    gui = SummaryGUI()
    gui.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()
