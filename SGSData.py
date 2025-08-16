# SGSData.py — one TABLE + two GRAPHS per sheet (last 6 months)
# Usage: SGSData.py <excel.xlsx> <months_csv> <year> [output.docx]

import sys, re
from pathlib import Path
from datetime import datetime, timedelta, date as ddate, time as dtime
from io import BytesIO

import openpyxl
from openpyxl.utils.datetime import from_excel as oxl_from_excel, CALENDAR_WINDOWS_1900

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import matplotlib.pyplot as plt
import matplotlib.dates as mdates

BASE_DIR    = Path(__file__).parent.resolve()
PRODUCT_DIR = BASE_DIR / "Product"; PRODUCT_DIR.mkdir(parents=True, exist_ok=True)

# -------- fixed chart display size (inches) --------
GRAPH_WIDTH_IN = 6.7   # phone-friendly fixed width

# -------- parsing config --------
DATE_PARSE_FORMATS = ("%d-%b-%Y","%d-%b-%y","%Y-%m-%d","%b %d, %Y","%B %d, %Y","%d/%m/%Y","%m/%d/%Y")
PARAM_KEYS  = {"cbod","cbod5","tss","tp","tan","tkn","no3","no2","tn","bod","bod5"}
GROUP1_KEYS = {"cbod5","bod5","tss","cbod","bod"}   # Graph 1
GROUP2_KEYS = {"tkn","tan","no2","no3","tn"}        # Graph 2
EXCLUDE_HDR_WORDS = {"units","objective","limit","average","median","cofa","eca"}

def find_latest_docx(product_dir: Path) -> Path:
    docs = [p for p in product_dir.glob("*.docx") if not p.name.startswith("~$")]
    if not docs: raise FileNotFoundError(f"No .docx files found in {product_dir}")
    return max(docs, key=lambda p: p.stat().st_mtime)

def ym_add(year: int, month: int, delta: int) -> tuple[int,int]:
    """Add delta months to (year, month)."""
    y = year + (month - 1 + delta) // 12
    m = (month - 1 + delta) % 12 + 1
    return y, m

def month_start(year: int, month: int) -> datetime:
    return datetime(year, month, 1)

def month_end(year: int, month: int) -> datetime:
    ny, nm = ym_add(year, month, 1)
    return datetime(ny, nm, 1) - timedelta(days=1)

def parse_date_cell(val) -> datetime | None:
    if isinstance(val, datetime): return val
    if val is None or (isinstance(val, str) and not val.strip()): return None
    if isinstance(val, str):
        s = val.strip()
        if s.lower() in ("units","average","median","cofa objective","eca objective","cofa limit","eca limit"):
            return None
        for fmt in DATE_PARSE_FORMATS:
            try: return datetime.strptime(s, fmt)
            except Exception: continue
        return None
    if isinstance(val, (int, float)):
        try: converted = oxl_from_excel(val, CALENDAR_WINDOWS_1900)
        except Exception:
            try: converted = datetime(1899,12,30) + timedelta(days=float(val))
            except Exception: return None
        if isinstance(converted, dtime): return None
        if isinstance(converted, ddate) and not isinstance(converted, datetime):
            return datetime(converted.year, converted.month, converted.day)
        return converted
    if isinstance(val, dtime): return None
    if isinstance(val, ddate): return datetime(val.year, val.month, val.day)
    return None

def text(ws, r, c) -> str:
    v = ws.cell(row=r, column=c).value
    return ("" if v is None else str(v)).strip()

def add_word_table(doc: Document, headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers)); t.style = "Table Grid"
    for j, h in enumerate(headers): t.cell(0, j).text = h
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row): t.cell(i, j).text = "" if v is None else str(v)
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                pf = p.paragraph_format; pf.space_before = Pt(0); pf.space_after = Pt(0)
    tblPr = t._tblPr; tblLayout = OxmlElement('w:tblLayout'); tblLayout.set(qn('w:type'), 'fixed'); tblPr.append(tblLayout)
    return t

# -------- find header/Date --------
def find_param_header_row(ws) -> int | None:
    max_r = min(ws.max_row, 80); max_c = min(ws.max_column, 80)
    for r in range(1, max_r+1):
        hits = 0
        for c in range(1, max_c+1):
            s = text(ws, r, c).lower()
            if not s: continue
            tok = re.sub(r"[^a-z0-9]+", "", s)
            if any(key in tok for key in PARAM_KEYS): hits += 1
        if hits >= 2: return r
    return None

def find_date_column(ws) -> int | None:
    max_r = min(ws.max_row, 60); max_c = min(ws.max_column, 60)
    # Prefer a header named "Date"
    for r in range(1, max_r+1):
        for c in range(1, max_c+1):
            if "date" in text(ws, r, c).lower(): return c
    # Fallback: best column by "looks like a date" density
    best_col, best_score = None, -1; N = min(ws.max_row, 40)
    for c in range(1, max_c+1):
        score = sum(1 for r in range(1, N+1) if parse_date_cell(ws.cell(row=r, column=c).value))
        if score > best_score: best_score, best_col = score, c
    return best_col

# -------- plotting (absolute, fixed size) --------
def plot_series_to_doc(doc: Document, title: str, series: dict[str, list[tuple[datetime,float]]]):
    if not series: return
    clean = {}
    for k, pts in series.items():
        pts = [(dt, v) for dt, v in pts if isinstance(dt, datetime)]
        pts.sort(key=lambda x: x[0])
        if pts: clean[k] = pts
    if not clean: return

    plt.figure(figsize=(11, 6), dpi=300)
    any_pts = False
    for label, pts in clean.items():
        xs = [dt for dt,_ in pts]; ys = [v for _,v in pts]
        if xs and ys:
            any_pts = True
            plt.plot(xs, ys, marker='o', linewidth=2.6, markersize=6, label=label)
    if not any_pts: plt.close(); return

    ax = plt.gca()
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%d-%b-%y'))
    ax.xaxis.set_major_locator(mdates.AutoDateLocator())
    plt.xticks(rotation=45, ha='right', fontsize=12)
    plt.yticks(fontsize=12)
    plt.ylabel("Concentration (mg/L)", fontsize=13)
    plt.title(title, fontsize=16)
    plt.legend(fontsize=12)
    plt.tight_layout()

    img = BytesIO(); plt.savefig(img, format="png", dpi=300); plt.close(); img.seek(0)
    doc.add_paragraph()
    doc.add_picture(img, width=Inches(GRAPH_WIDTH_IN))
    img.close()

# -------- per-sheet workflow (one table + two graphs over last 6 months) --------
def table_then_two_graphs(doc: Document, ws, sheet_name: str, months_csv: str, year: int, need_pagebreak: bool) -> bool:
    param_row = find_param_header_row(ws); date_col = find_date_column(ws)
    if not param_row or not date_col: return False

    # Data begins after the param row; skip meta rows like Units/ECA/CofA
    start_row = param_row + 1
    for _ in range(2):
        row_text = " ".join(text(ws, start_row, c).lower() for c in range(1, ws.max_column+1))
        if any(k in row_text for k in ("units","cofa","eca","objective","limit")): start_row += 1

    # Collect all dates + row indices
    dates, row_idxs = [], []
    for r in range(start_row, ws.max_row+1):
        dt = parse_date_cell(ws.cell(row=r, column=date_col).value)
        if dt: dates.append(dt); row_idxs.append(r)
    if not dates: return False

    # --- TRUE last-6-months window (cross-year safe) ---
    tokens = [int(x) for x in months_csv.split(",") if x.strip().isdigit()]
    if tokens:
        sel_month = tokens[-1]     # GUI's selected month is last in months_csv (AutoSummary behavior)
        sel_year  = year
    else:
        latest = max(dates)
        sel_month, sel_year = latest.month, latest.year

    start_y, start_m = ym_add(sel_year, sel_month, -5)
    start_dt = month_start(start_y, start_m)
    end_dt   = month_end(sel_year, sel_month)

    idxs = [i for i, d in enumerate(dates) if start_dt <= d <= end_dt]
    if not idxs: return False

    # Select parameter columns & groups (only those with any numeric data in the 6-month window)
    group1_cols, group1_labels, group2_cols, group2_labels = [], [], [], []
    all_cols, all_labels = [], []
    for c in range(1, ws.max_column+1):
        if c == date_col: continue
        raw = text(ws, param_row, c)
        if not raw: continue
        norm = re.sub(r"[^a-z0-9]+", "", raw.lower())
        if not norm or norm == "date" or any(w in norm for w in EXCLUDE_HDR_WORDS): continue

        in_g1 = any(k in norm for k in GROUP1_KEYS)
        in_g2 = any(k in norm for k in GROUP2_KEYS)
        if not (in_g1 or in_g2): continue

        has_num = any(isinstance(ws.cell(row=row_idxs[i], column=c).value, (int,float)) for i in idxs)
        if not has_num: continue

        all_cols.append(c); all_labels.append(raw.strip())
        if in_g1: group1_cols.append(c); group1_labels.append(raw.strip())
        if in_g2: group2_cols.append(c); group2_labels.append(raw.strip())
    if not all_cols: return False

    # PAGE BREAK per sheet (not per month)
    if need_pagebreak:
        doc.add_page_break()

    # Heading uses trimmed sheet name (drop common 3-letter site code prefix if present)
    trimmed = sheet_name.split(" ", 1)[1] if " " in sheet_name else sheet_name
    doc.add_heading(trimmed, level=2)

    # TABLE over the full 6-month window
    headers = ["Date"] + all_labels
    rows = []
    for i in idxs:
        r = row_idxs[i]; dt = dates[i]
        row = [dt.strftime("%d-%b-%y")]
        for c in all_cols:
            v = ws.cell(row=r, column=c).value
            try: row.append(f"{float(v):g}")
            except Exception: row.append("" if v is None else str(v))
        rows.append(row)

    add_word_table(doc, headers, rows)
    doc.add_paragraph()

    # GRAPH 1: cBOD/BOD/TSS — last 6 months (all points in window)
    if group1_cols:
        series1 = {}
        for label, c in zip(group1_labels, group1_cols):
            pts = []
            for i in idxs:
                r = row_idxs[i]; v = ws.cell(row=r, column=c).value
                try: pts.append((dates[i], float(v)))
                except Exception: continue
            if pts: series1[label] = pts
        if series1: plot_series_to_doc(doc, f"{trimmed} — cBOD/BOD/TSS (Last 6 Months)", series1)

    # GRAPH 2: Nitrogen species — last 6 months
    if group2_cols:
        series2 = {}
        for label, c in zip(group2_labels, group2_cols):
            pts = []
            for i in idxs:
                r = row_idxs[i]; v = ws.cell(row=r, column=c).value
                try: pts.append((dates[i], float(v)))
                except Exception: continue
            if pts: series2[label] = pts
        if series2: plot_series_to_doc(doc, f"{trimmed} — Nitrogen Species (Last 6 Months)", series2)

    return True

def main():
    if len(sys.argv) < 4:
        print("Usage: SGSData.py <excel.xlsx> <months_csv> <year> [output.docx]"); return
    excel_path = Path(sys.argv[1]); months_csv = sys.argv[2]; year = int(sys.argv[3])
    out_docx = Path(sys.argv[4]) if len(sys.argv) >= 5 else None

    wb = openpyxl.load_workbook(excel_path, data_only=True)
    doc_path = out_docx if out_docx else find_latest_docx(PRODUCT_DIR)
    doc = Document(doc_path)

    # Remove top margin across sections
    for section in doc.sections:
        section.top_margin = Pt(0)

    # Keep your appendix header
    doc.add_page_break()
    doc.add_heading("Appendix A: SGS Tables & Graphs", level=1)

    any_done = False
    first_section = True
    for sheet_name in wb.sheetnames:
        lname = sheet_name.lower()
        process = (any(k in lname for k in ["raw","sewage","biofilter","waternox","waternox-ls"])
                   or ("final effluent" in lname or "polisher effluent" in lname))
        if not process: continue

        if table_then_two_graphs(doc, wb[sheet_name], sheet_name, months_csv, year, need_pagebreak=not first_section):
            any_done = True
            first_section = False

        # Stop after final/polisher (matches original flow)
        if "final effluent" in lname or "polisher effluent" in lname:
            break

    doc.save(doc_path)
    print(f"{'Appended' if any_done else 'No'} SGS tables & graphs {'into' if any_done else ''} {doc_path.name if any_done else ''}".strip())

if __name__ == "__main__": main()
