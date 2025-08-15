# FlowData.py
import sys
import re
from pathlib import Path
from datetime import datetime
from io import BytesIO
import matplotlib.pyplot as plt

from docx import Document
import openpyxl
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.enum.table import WD_TABLE_ALIGNMENT
from collections import defaultdict

# ---------- Helper: insert paragraph after ----------
def insert_paragraph_after(paragraph: Paragraph, text: str, style: str = None) -> Paragraph:
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    new_para.add_run(text)
    return new_para

# ---------- paths ----------
BASE_DIR    = Path(__file__).parent.resolve()
PRODUCT_DIR = BASE_DIR / "Product"
PRODUCT_DIR.mkdir(parents=True, exist_ok=True)

LIGHT_BLUE_HEX = "D9EAF7"

def find_latest_docx(product_dir: Path) -> Path:
    docs = [p for p in product_dir.glob("*.docx") if not p.name.startswith("~$")]
    if not docs:
        raise FileNotFoundError(f"No usable .docx files found in {product_dir}")
    return max(docs, key=lambda p: p.stat().st_mtime)

def set_cell_width(cell, width_cm: float):
    width = Cm(width_cm)
    tcPr  = cell._tc.get_or_add_tcPr()
    for w in tcPr.findall(qn('w:tcW')):
        tcPr.remove(w)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    cell.width = width

def set_table_column_widths(table):
    table.allow_autofit = False
    for i, col in enumerate(table.columns):
        col_w = 4 if i == 0 else 2.5
        for cell in col.cells:
            set_cell_width(cell, col_w)

def shade_row_hex(row, hex_colour: str):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),  'clear')
        shd.set(qn('w:fill'), hex_colour)
        tcPr.append(shd)
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.line_spacing = 1

def extract_peak_capacity(doc: Document) -> float | None:
    pattern = re.compile(
        r'(?:peak rated capacity|peak design daily flow rated capacity of)\s*.*?([\d,]+(?:\.\d+)?)\s*L/day',
        re.IGNORECASE
    )
    for para in doc.paragraphs:
        match = pattern.search(para.text)
        if match:
            value_str = match.group(1).replace(',', '')
            try:
                return float(value_str)
            except ValueError:
                continue
    return None

def sheet_to_word_table(ws, title: str, doc, site: str, peak_capacity: float | None = None):
    # locate Date row and Average row
    from datetime import datetime as _dt
    date_row = avg_row = None
    for r in range(1, ws.max_row + 1):
        first_val = ws.cell(row=r, column=1).value
        if isinstance(first_val, str) and first_val.strip().lower() == "date":
            date_row = r
        row_text = " ".join(
            str(ws.cell(row=r, column=c).value).lower()
            for c in range(1, ws.max_column + 1)
            if ws.cell(row=r, column=c).value is not None
        )
        if date_row and "average" in row_text:
            avg_row = r
            break
    if not date_row or not avg_row or avg_row < date_row:
        print(f"→ Sheet '{ws.title}': missing 'Date' or 'Average' rows.")
        return

    rows = range(date_row, avg_row + 1)
    import re as _re
    pump_pat = _re.compile(r'^pump\s*\d', _re.IGNORECASE)
    keep_cols = []
    for c in range(1, ws.max_column + 1):
        raw = ws.cell(row=date_row, column=c).value
        header = str(raw).strip().lower() if raw is not None else ''
        if 'daily' in header and not header.startswith('total daily flow'):
            break
        if pump_pat.search(header):
            continue
        if any(ws.cell(row=r_, column=c).value not in (None, "") for r_ in rows):
            keep_cols.append(c - 1)
    if not keep_cols:
        print(f"→ Sheet '{ws.title}': no data columns before or including 'Daily'.")
        return

    doc.add_heading(title, level=2)
    table = doc.add_table(rows=len(rows), cols=len(keep_cols))
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    chart_col_idx = None

    for ridx, r in enumerate(rows):
        for cidx, col_idx in enumerate(keep_cols):
            val = ws.cell(row=r, column=col_idx + 1).value
            text = (
                val.strftime("%d-%b-%y") if isinstance(val, _dt)
                else (f"{val:.0f}" if isinstance(val, (float, int)) else (str(val) if val is not None else ""))
            )
            cell = table.cell(ridx, cidx)
            cell.text = text
            if r == date_row and isinstance(val, str) and val.strip().lower() == "flow":
                chart_col_idx = cidx

    set_table_column_widths(table)
    shade_row_hex(table.rows[0],  LIGHT_BLUE_HEX)
    shade_row_hex(table.rows[-1], LIGHT_BLUE_HEX)
    shade_row_hex(table.rows[-2], LIGHT_BLUE_HEX)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                pf = paragraph.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after  = Pt(0)
                pf.line_spacing = 1

    # find Flow column
    excel_chart_col = None
    for col_idx in keep_cols:
        header_val = ws.cell(row=date_row, column=col_idx + 1).value
        if isinstance(header_val, str) and header_val.strip().lower() == "flow":
            excel_chart_col = col_idx + 1
            break
    if not excel_chart_col:
        excel_chart_col = keep_cols[-1] + 1
        chart_col_idx = len(keep_cols) - 1

    # --- series ---
    times, values = [], []
    exceedances = []
    for ridx, r in enumerate(rows):
        date_val = ws.cell(row=r, column=1).value
        dt = None
        if isinstance(date_val, _dt):
            dt = date_val
        elif date_val:
            for fmt in ("%d-%b-%y", "%d-%b-%Y", "%Y-%m-%d", "%b %d, %Y", "%B %d, %Y"):
                try:
                    dt = _dt.strptime(str(date_val), fmt)
                    break
                except Exception:
                    continue
        val = ws.cell(row=r, column=excel_chart_col).value
        try:
            fval = float(val)
        except Exception:
            continue
        times.append(dt)
        values.append(fval)
        if peak_capacity is not None and fval > peak_capacity and chart_col_idx is not None:
            exceedances.append((dt, fval))
            cell = table.cell(ridx, chart_col_idx)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    # build by-day chart (fallback if dates bad)
    if values and all(t is None for t in times):
        days = list(range(1, len(values)+1))
    else:
        pairs = [(t, v) for t, v in zip(times, values) if isinstance(t, datetime)]
        days = [t.day for t, _ in pairs]
        values = [v for _, v in pairs]

    if values and days:
        plt.figure(figsize=(8, 4))
        plt.plot(days, values, marker='o')
        if peak_capacity is not None:
            try:
                plt.hlines(peak_capacity, min(days), max(days), linestyles='--')
            except ValueError:
                pass
            if exceedances:
                xs = [dt.day for dt, _ in exceedances if isinstance(dt, datetime)]
                ys = [v for dt, v in exceedances if isinstance(dt, datetime)]
                if xs and ys:
                    plt.scatter(xs, ys)

        # title: expand month abbrev to full if possible
        month_abbrev = title.split()[0]
        try:
            month_full = datetime.strptime(month_abbrev, '%b').strftime('%B')
        except Exception:
            month_full = month_abbrev
        plt.title(f"{month_full} flow data of {site}")
        plt.xticks(days, rotation=45, ha='right')
        plt.xlabel(f"Day in {month_full}")
        plt.ylabel("Flow")
        plt.tight_layout()
        img_stream = BytesIO()
        plt.savefig(img_stream, format="png")
        plt.close()
        img_stream.seek(0)
        doc.add_paragraph()
        doc.add_picture(img_stream, width=Cm(16))
        img_stream.close()

        if peak_capacity is not None and values:
            exceed_count = len([v for _, v in exceedances])
            avg_flow = sum(values) / len(values)
            avg_classification = "well within" if avg_flow < 0.9 * peak_capacity else ("close to" if avg_flow <= peak_capacity else "above")
            summary_text = (
                f"{exceed_count} day(s) exceeded the peak rated capacity of {int(peak_capacity):,} L/day. "
                f"The average daily flow remained {avg_classification} the anticipated range."
            )
            heading_pat = re.compile(
                r'flow\s+discharged\s+to\s+(?:the\s+)?subsurface\s+(?:disposal|dispersal)\s+system',
                re.IGNORECASE
            )
            insertion_para = None
            for para in doc.paragraphs:
                if heading_pat.fullmatch(para.text.strip()):
                    insertion_para = para
                    break
            if insertion_para:
                insert_paragraph_after(insertion_para, summary_text)
            else:
                doc.add_paragraph(summary_text)

def main():
    # CLI:
    # 1) <excel_path>
    # 2) <months_csv> like "3,4,5"
    # 3) <year>
    # 4) [optional] <output_docx_path>
    if len(sys.argv) < 4:
        print("Usage: FlowData.py <excel.xlsx> <months_csv> <year> [output.docx]")
        return

    excel_path = Path(sys.argv[1])
    months     = [int(x) for x in sys.argv[2].split(",") if x.strip()]
    year       = int(sys.argv[3])
    out_docx   = Path(sys.argv[4]) if len(sys.argv) >= 5 else None

    wb = openpyxl.load_workbook(excel_path, data_only=True)
    site = excel_path.stem

    if out_docx:
        doc_path = out_docx
    else:
        doc_path = find_latest_docx(PRODUCT_DIR)

    doc = Document(doc_path)
    peak_capacity = extract_peak_capacity(doc)

    month_names_full = ["January","February","March","April","May","June",
                        "July","August","September","October","November","December"]

    for m in sorted(months):
        month_name     = month_names_full[m-1]
        two_digit_year = str(year)[-2:]
        target         = f"{month_name[:3].title()} {two_digit_year}"

        matching_sheets = [s for s in wb.sheetnames if target.lower() in s.lower()]
        if not matching_sheets:
            print(f"No sheet containing '{target}' found in {excel_path.name}")
            continue

        doc.add_page_break()
        doc.add_heading(f"Appendix B: Flow Data — {month_name} {year}", level=1)

        made_chart = False
        before_shapes = len(doc.inline_shapes)
        for sheet_name in matching_sheets:
            sheet_to_word_table(wb[sheet_name], sheet_name, doc, site, peak_capacity=peak_capacity)
        made_chart = len(doc.inline_shapes) > before_shapes

        if not made_chart:
            # Minimal by-index fallback
            ws = wb[matching_sheets[-1]]
            date_row = None
            for r in range(1, ws.max_row + 1):
                if str(ws.cell(row=r, column=1).value).strip().lower() == "date":
                    date_row = r; break
            if date_row:
                cols = []
                for c in range(2, ws.max_column + 1):
                    any_num = any(isinstance(ws.cell(row=rr, column=c).value, (int,float))
                                  for rr in range(date_row+1, ws.max_row+1))
                    if any_num: cols.append(c)
                if cols:
                    vals = []
                    for rr in range(date_row+1, ws.max_row+1):
                        v = ws.cell(row=rr, column=cols[-1]).value
                        if isinstance(v, (int,float)): vals.append(v)
                    if vals:
                        img = BytesIO()
                        plt.figure(figsize=(8,4))
                        plt.plot(range(1, len(vals)+1), vals, marker='o')
                        plt.title(f"{month_name} flow data of {site}")
                        plt.xlabel(f"Day in {month_name}")
                        plt.ylabel("Flow")
                        plt.tight_layout()
                        plt.savefig(img, format="png"); plt.close(); img.seek(0)
                        doc.add_paragraph()
                        doc.add_picture(img, width=Cm(16))
                        img.close()

        doc.add_heading("Appendix C: Site Notes", level=1)

    doc.save(doc_path)
    print(f"Appended FlowData results into {doc_path.name}")

if __name__ == "__main__":
    main()
